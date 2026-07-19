"""Parser tests (offline). DOCX/XLSX fixtures are generated in-test."""

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

from hermes_assistant.rag.parsers import (
    SUPPORTED_SUFFIXES,
    UnsupportedFileError,
    parse_file,
)


def test_parse_markdown_is_verbatim(sample_md: Path) -> None:
    """Markdown is read through unchanged, headings intact."""
    text = parse_file(sample_md)
    assert "# HERMES 2022 Overview" in text
    assert "## Phases" in text
    assert "Initialisierung" in text


def test_parse_txt(tmp_path: Path) -> None:
    """Plain .txt files are read as Markdown/text."""
    p = tmp_path / "note.txt"
    p.write_text("hello world", encoding="utf-8")
    assert parse_file(p) == "hello world"


def test_parse_docx_maps_headings(tmp_path: Path) -> None:
    """DOCX Heading styles become Markdown headings."""
    doc = Document()
    doc.add_heading("Phases", level=1)
    doc.add_paragraph("Initialisierung and Konzept.")
    doc.add_heading("Roles", level=2)
    doc.add_paragraph("Projektleiter steers the work.")
    path = tmp_path / "d.docx"
    doc.save(str(path))

    text = parse_file(path)
    assert "# Phases" in text
    assert "## Roles" in text
    assert "Projektleiter steers the work." in text


def test_parse_xlsx_sheet_becomes_heading(tmp_path: Path) -> None:
    """XLSX sheet names become headings; rows are tab-joined."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Outcomes"
    ws.append(["Projektauftrag", "mandatory"])
    ws.append(["Studie", "mandatory"])
    path = tmp_path / "s.xlsx"
    wb.save(str(path))

    text = parse_file(path)
    assert "# Outcomes" in text
    assert "Projektauftrag\tmandatory" in text


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    """An unknown extension raises UnsupportedFileError."""
    p = tmp_path / "x.rtf"
    p.write_text("x", encoding="utf-8")
    with pytest.raises(UnsupportedFileError):
        parse_file(p)


def test_missing_file_raises(tmp_path: Path) -> None:
    """A non-existent path raises FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        parse_file(tmp_path / "nope.md")


def test_supported_suffixes_cover_expected() -> None:
    """The supported-suffix set covers the four target formats."""
    assert {".pdf", ".docx", ".xlsx", ".md"} <= SUPPORTED_SUFFIXES
